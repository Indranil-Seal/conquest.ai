"""
conquest.ai — Document Indexer

Handles loading, parsing, chunking, and embedding of documents from the DREAM library.
Supports PDF, EPUB, and DOCX formats.

Workflow:
    1. Scan DREAM/library/ for supported file types
    2. Parse each file into raw text using the appropriate loader
    3. Split text into chunks with LlamaIndex SentenceSplitter
    4. Embed chunks with sentence-transformers (all-MiniLM-L6-v2) — runs locally
    5. Store embeddings + metadata in ChromaDB
"""

import os
import logging
from pathlib import Path
from typing import Optional

import chromadb
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

from llama_index.core import Document, VectorStoreIndex, StorageContext, Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore

logger = logging.getLogger(__name__)

# --- Configuration ---
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
CHUNK_SIZE = 512      # tokens per chunk — good balance for technical text
CHUNK_OVERLAP = 50    # token overlap between chunks to preserve context
CHROMA_COLLECTION = "conquest_dream"
SUPPORTED_EXTENSIONS = {".pdf", ".epub", ".docx"}


def load_pdf(file_path: Path) -> str:
    """
    Extract text from a PDF file using PyMuPDF.

    PyMuPDF is chosen over pypdf because it handles mathematical symbols,
    subscripts, and superscripts more accurately — critical for DS/ML papers.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Extracted text as a single string (pages joined with newlines).
    """
    text_parts = []
    try:
        with fitz.open(str(file_path)) as doc:
            for page in doc:
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(text)
    except Exception as e:
        logger.warning(f"Failed to parse PDF {file_path.name}: {e}")
    return "\n".join(text_parts)


def load_epub(file_path: Path) -> str:
    """
    Extract text from an EPUB file using ebooklib + BeautifulSoup.

    ebooklib reads the EPUB container; BeautifulSoup strips HTML tags
    from each chapter's XHTML content.

    Args:
        file_path: Path to the EPUB file.

    Returns:
        Extracted text as a single string.
    """
    try:
        import ebooklib
        from ebooklib import epub

        book = epub.read_epub(str(file_path))
        text_parts = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "html.parser")
            text = soup.get_text(separator="\n", strip=True)
            if text.strip():
                text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.warning(f"Failed to parse EPUB {file_path.name}: {e}")
        return ""


def load_docx(file_path: Path) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Reads each paragraph in document order and joins them with newlines.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text as a single string.
    """
    try:
        doc = DocxDocument(str(file_path))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        logger.warning(f"Failed to parse DOCX {file_path.name}: {e}")
        return ""


def classify_document(filename: str) -> str:
    """
    Heuristically classify a document as 'textbook', 'paper', or 'reference'
    based on its filename. Used as metadata for better citation formatting.

    Args:
        filename: The filename (without path) of the document.

    Returns:
        Document type string.
    """
    name_lower = filename.lower()
    if any(k in name_lower for k in ["introduction", "hands-on", "fundamentals", "beginning",
                                      "elements", "statistical learning", "programming"]):
        return "textbook"
    if any(k in name_lower for k in ["cheat", "reference", "cheatsheet", "guide", "quick"]):
        return "reference"
    return "paper"


def load_documents_from_dream(dream_path: Path) -> list[Document]:
    """
    Scan the DREAM library directory and load all supported documents.

    For each file found:
    - Extracts text using the appropriate parser
    - Wraps the text in a LlamaIndex Document with metadata
    - Skips empty or unparseable files with a warning

    Args:
        dream_path: Path to the root of the DREAM repository.

    Returns:
        List of LlamaIndex Document objects ready for indexing.
    """
    library_path = dream_path / "library"
    if not library_path.exists():
        # Fall back to root if no library/ subdirectory
        library_path = dream_path

    documents = []
    files = [f for f in library_path.rglob("*") if f.suffix.lower() in SUPPORTED_EXTENSIONS]

    if not files:
        logger.warning(f"No supported files found in {library_path}")
        return documents

    logger.info(f"Found {len(files)} document(s) in DREAM library")

    for file_path in files:
        ext = file_path.suffix.lower()
        logger.info(f"  Loading: {file_path.name}")

        if ext == ".pdf":
            text = load_pdf(file_path)
        elif ext == ".epub":
            text = load_epub(file_path)
        elif ext == ".docx":
            text = load_docx(file_path)
        else:
            continue

        if not text.strip():
            logger.warning(f"  Skipping {file_path.name} — no text extracted")
            continue

        doc = Document(
            text=text,
            metadata={
                "filename": file_path.name,
                "file_type": ext.lstrip("."),
                "doc_type": classify_document(file_path.name),
                "source": str(file_path.relative_to(dream_path)),
            },
            metadata_seperator="\n",
            metadata_template="{key}: {value}",
            text_template="File: {metadata_str}\n\n{content}",
        )
        documents.append(doc)

    logger.info(f"Loaded {len(documents)} document(s) successfully")
    return documents


def build_index(
    dream_path: Path,
    chroma_db_path: Path,
    force_reindex: bool = False,
) -> VectorStoreIndex:
    """
    Build (or load) the ChromaDB vector index from DREAM library documents.

    If the ChromaDB collection already contains documents and force_reindex
    is False, loading the existing index is skipped (fast startup).

    If the collection is empty or force_reindex=True, documents are loaded,
    chunked, embedded, and stored in ChromaDB.

    Args:
        dream_path: Path to the DREAM repository root.
        chroma_db_path: Path to the ChromaDB persistent storage directory.
        force_reindex: If True, re-index all documents even if index exists.

    Returns:
        A LlamaIndex VectorStoreIndex backed by ChromaDB.
    """
    # Configure embedding model (runs locally — no API calls)
    embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)
    Settings.embed_model = embed_model
    # Disable the default LlamaIndex LLM — we use Claude via the RAG pipeline
    Settings.llm = None

    # Set up ChromaDB
    chroma_client = chromadb.PersistentClient(path=str(chroma_db_path))
    chroma_collection = chroma_client.get_or_create_collection(CHROMA_COLLECTION)
    vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    # Check if the collection already has documents
    existing_count = chroma_collection.count()
    if existing_count > 0 and not force_reindex:
        logger.info(
            f"ChromaDB collection '{CHROMA_COLLECTION}' already has {existing_count} chunks. "
            "Loading existing index (use --force to re-index)."
        )
        return VectorStoreIndex.from_vector_store(
            vector_store=vector_store,
            storage_context=storage_context,
        )

    # Load and index documents
    logger.info("Building index from DREAM library...")
    documents = load_documents_from_dream(dream_path)

    if not documents:
        raise RuntimeError(
            f"No documents loaded from DREAM library at {dream_path}. "
            "Please ensure the DREAM repository is cloned and contains files in library/."
        )

    # Chunk documents with SentenceSplitter
    splitter = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)

    index = VectorStoreIndex.from_documents(
        documents,
        storage_context=storage_context,
        transformations=[splitter],
        show_progress=True,
    )

    total_chunks = chroma_collection.count()
    logger.info(f"Indexing complete. {total_chunks} chunks stored in ChromaDB.")
    return index


def remove_documents_from_index(
    chroma_db_path: Path,
    filenames: list[str],
) -> int:
    """
    Remove all chunks belonging to the given filenames from ChromaDB.

    Called when ingest.py detects that documents were deleted or modified
    in the DREAM library. Modified documents are first removed here, then
    re-indexed with fresh content by index_specific_documents().

    Args:
        chroma_db_path: Path to the ChromaDB persistent storage directory.
        filenames: List of filenames (basename only) whose chunks should be deleted.

    Returns:
        Total number of chunks removed.
    """
    if not filenames:
        return 0

    chroma_client = chromadb.PersistentClient(path=str(chroma_db_path))
    chroma_collection = chroma_client.get_or_create_collection(CHROMA_COLLECTION)

    removed_count = 0
    for filename in filenames:
        # Retrieve IDs of all chunks that belong to this file
        results = chroma_collection.get(
            where={"filename": filename},
            include=[],  # only IDs needed
        )
        if results["ids"]:
            chroma_collection.delete(ids=results["ids"])
            removed_count += len(results["ids"])
            logger.info(f"  Removed {len(results['ids'])} chunks for: {filename}")
        else:
            logger.info(f"  No chunks found for: {filename} (skipping)")

    return removed_count


def index_specific_documents(
    dream_path: Path,
    chroma_db_path: Path,
    file_paths: list[Path],
) -> int:
    """
    Index only the specified document files into the existing ChromaDB collection.

    Used for incremental updates when ingest.py detects new or modified files
    in the DREAM library after a git pull. The collection is not cleared —
    new chunks are appended alongside existing ones.

    Args:
        dream_path: Root of the DREAM repository (used to compute relative paths).
        chroma_db_path: Path to the ChromaDB persistent storage directory.
        file_paths: Absolute paths to the specific files to index.

    Returns:
        Number of documents successfully indexed.
    """
    if not file_paths:
        return 0

    embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)
    Settings.embed_model = embed_model
    Settings.llm = None

    chroma_client = chromadb.PersistentClient(path=str(chroma_db_path))
    chroma_collection = chroma_client.get_or_create_collection(CHROMA_COLLECTION)
    vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    documents = []
    for file_path in file_paths:
        if not file_path.exists():
            logger.warning(f"  File not found (skipping): {file_path.name}")
            continue

        ext = file_path.suffix.lower()
        logger.info(f"  Indexing: {file_path.name}")

        if ext == ".pdf":
            text = load_pdf(file_path)
        elif ext == ".epub":
            text = load_epub(file_path)
        elif ext == ".docx":
            text = load_docx(file_path)
        else:
            continue

        if not text.strip():
            logger.warning(f"  Skipping {file_path.name} — no text extracted")
            continue

        documents.append(Document(
            text=text,
            metadata={
                "filename": file_path.name,
                "file_type": ext.lstrip("."),
                "doc_type": classify_document(file_path.name),
                "source": str(file_path.relative_to(dream_path)),
            },
            metadata_seperator="\n",
            metadata_template="{key}: {value}",
            text_template="File: {metadata_str}\n\n{content}",
        ))

    if not documents:
        logger.warning("No documents were successfully loaded for incremental indexing.")
        return 0

    splitter = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    VectorStoreIndex.from_documents(
        documents,
        storage_context=storage_context,
        transformations=[splitter],
        show_progress=True,
    )

    logger.info(f"Incremental indexing complete. {len(documents)} document(s) added.")
    return len(documents)


def load_existing_index(chroma_db_path: Path) -> Optional[VectorStoreIndex]:
    """
    Load a pre-built ChromaDB index without re-indexing documents.

    Used by the Chainlit app at startup for fast initialization.

    Args:
        chroma_db_path: Path to the ChromaDB persistent storage directory.

    Returns:
        VectorStoreIndex if the collection exists and has data, else None.
    """
    embed_model = HuggingFaceEmbedding(model_name=EMBEDDING_MODEL)
    Settings.embed_model = embed_model
    Settings.llm = None

    try:
        chroma_client = chromadb.PersistentClient(path=str(chroma_db_path))
        chroma_collection = chroma_client.get_or_create_collection(CHROMA_COLLECTION)

        if chroma_collection.count() == 0:
            return None

        vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
        storage_context = StorageContext.from_defaults(vector_store=vector_store)

        return VectorStoreIndex.from_vector_store(
            vector_store=vector_store,
            storage_context=storage_context,
        )
    except Exception as e:
        logger.error(f"Failed to load ChromaDB index: {e}")
        return None
